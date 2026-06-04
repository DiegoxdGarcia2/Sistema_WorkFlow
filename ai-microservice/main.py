"""
═══════════════════════════════════════════════════════════════════
  BPM Inteligente — AI Microservice (FastAPI)
  Fase 2: Groq Real + Streaming SSE + RAG Contextual
═══════════════════════════════════════════════════════════════════
"""

import os
import json
import asyncio
import io
from fastapi import FastAPI, HTTPException, UploadFile, File, Request, Form
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from typing import Optional, List, Dict
from groq import Groq
from elevenlabs import ElevenLabs
from dotenv import load_dotenv

# ── Cargar variables de entorno ──────────────────────────────────
load_dotenv(override=True)

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
ELEVENLABS_API_KEY = os.getenv("ELEVENLABS_API_KEY", "")
ELEVENLABS_VOICE_ID = os.getenv("ELEVENLABS_VOICE_ID", "EXAVITQu4vr4xnSDxMaL")
API_SECRET = os.getenv("AI_API_SECRET", "dev_secret_local_only")

# ── Inicializar clientes ────────────────────────────────────────
app = FastAPI(
    title="BPM Inteligente — AI Microservice",
    description="Microservicio de IA/ML con Groq, Streaming SSE y RAG contextual.",
    version="2.0.0",
)

groq_client: Groq | None = None
if GROQ_API_KEY:
    groq_client = Groq(api_key=GROQ_API_KEY)

el_client: ElevenLabs | None = None
if ELEVENLABS_API_KEY and ELEVENLABS_API_KEY != "your_key_here":
    el_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)

# ── CORS ─────────────────────────────────────────────────────────
_cors_origins = [
    "http://localhost:4200",
    "http://localhost:8080",
]
# Agregar orígenes de producción si están configurados
for env_key in ("CORS_ORIGIN", "BACKEND_ORIGIN"):
    val = os.getenv(env_key, "")
    if val:
        _cors_origins.append(val)

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_origin_regex=r"https://.*\.run\.app",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Seguridad inter-servicios ────────────────────────────────────
@app.middleware("http")
async def validate_api_secret(request: Request, call_next):
    # Permitir health check y CORS preflight sin secreto
    if request.url.path == "/" or request.method == "OPTIONS":
        return await call_next(request)

    # En desarrollo local, si el secreto es el default, permitir todo
    if API_SECRET == "dev_secret_local_only":
        return await call_next(request)

    # En producción, validar el header X-API-Secret
    secret = request.headers.get("X-API-Secret", "")
    if secret != API_SECRET:
        return JSONResponse(status_code=403, content={"detail": "Forbidden: invalid API secret"})

    return await call_next(request)


# ═══════════════════════════════════════════════════════════════════
# MODELOS PYDANTIC
# ═══════════════════════════════════════════════════════════════════

# ── Asistente IA (Designer) ──────────────────────────────────────
class AiCommandRequest(BaseModel):
    politicaId: Optional[str] = None
    instruccion: str
    contexto: Optional[dict] = None


class AiAction(BaseModel):
    tipo: str
    params: dict = Field(default_factory=dict)


class AiActionResponse(BaseModel):
    explicacion: str
    acciones: list[AiAction] = Field(default_factory=list)


# ── Chatbot ──────────────────────────────────────────────────────
class ChatMessage(BaseModel):
    role: str  # "user" | "assistant"
    content: str


class ChatbotRequest(BaseModel):
    mensaje: str
    contextoSeccion: Optional[str] = None
    contextoDinamico: Optional[str] = None
    historial: list[ChatMessage] = Field(default_factory=list)


class ChatbotResponse(BaseModel):
    respuesta: str
    rutaNavegacion: Optional[str] = None


from ml_service import ejecutar_pipeline_completo

# ── ML Analysis ──────────────────────────────────────────────────
class Finding(BaseModel):
    type: str
    severity: str
    nodeId: Optional[str] = ""
    message: str
    suggestion: str


class AnalysisRequest(BaseModel):
    politicaId: str
    registrosCompletados: list[dict] = Field(default_factory=list)
    tenantId: Optional[str] = None


class AnalysisResponse(BaseModel):
    findings: list[Finding] = Field(default_factory=list)

# ── Insights (Fase 3) ────────────────────────────────────────────
class BottleneckInfo(BaseModel):
    actividadId: str
    actividadNombre: str
    promedioMinutos: float
    desviacionSobre: float
    severity: str
    numEjecuciones: int

class Prediccion(BaseModel):
    duracionEstimadaDias: float
    confianza: float
    factoresRelevantes: list[str]

class Metrica(BaseModel):
    totalRegistros: int
    duracionPromedioMinutos: float
    desviacionEstandar: float
    tasaCompletitud: float

class InsightsResponse(BaseModel):
    politicaId: Optional[str]
    generadoEn: str
    metricas: Metrica
    cuellosBottella: list[BottleneckInfo]
    prediccion: Prediccion
    insightsNaturales: str
    alertas: list[dict]

# ── Móvil AI (Fase 3.2) ──────────────────────────────────────────
class VoiceRouterResponse(BaseModel):
    politicaId: str
    intencionDetectada: str
    confianza: float

# ═══════════════════════════════════════════════════════════════════
# SYSTEM PROMPTS
# ═══════════════════════════════════════════════════════════════════

CHATBOT_SYSTEM_PROMPT = """Eres BPM-Guía, el asistente experto y navegador del sistema BPM Inteligente.
Tu objetivo es ayudar al usuario y, si es posible, LLEVARLO al lugar correcto del sistema.

REGLAS DE RESPUESTA:
1. Responde de forma concisa, clara y en español. Usa Markdown ligero (negritas con **, listas con *).
2. Sé específico: di EXACTAMENTE qué botón presionar y dónde.
3. Usa los datos del sistema cuando estén disponibles para personalizar la respuesta.
4. NO repitas la misma respuesta genérica. Adapta según el contexto.
5. VALIDA EL ROL DEL USUARIO: En el contexto se te proveerá el rol del usuario (ej: Funcionario, Administrador, Diseñador). Si el usuario pide ir al diseñador de políticas o a la configuración, pero su rol no tiene los permisos lógicos, indícale amablemente que no tiene autorización y NO emitas la ruta de navegación.

RUTAS DE NAVEGACIÓN DISPONIBLES (Úsalas exactamente así):
- '/admin?tab=analytics': Dashboard de Analytics ML, predicciones de IA y cuellos de botella.
- '/admin?tab=monitor': Monitor de procesos en tiempo real.
- '/admin?tab=usuarios': Gestión de Colaboradores/Usuarios internos.
- '/admin?tab=clientes': Gestión de Clientes externos (NO es lo mismo que usuarios).
- '/admin?tab=departamentos': Estructura de Departamentos.
- '/admin?tab=cargos': Gestión de Cargos institucionales.
- '/admin?tab=tenants': Datos de la Empresa/Tenant.
- '/admin?tab=audit': Auditoría del sistema.
- '/admin?tab=formularios': Repositorio de Formularios.
- '/designer': Hub de Proyectos y nuevas políticas (Requiere rol Diseñador o Admin).
- '/designer/editor': Editor gráfico de flujos.
- '/funcionario?tab=bandeja': Bandeja de tareas pendientes del usuario.
- '/funcionario?tab=disponible': Mercado de tareas disponibles para tomar.
- '/funcionario?tab=historial': Mi Historial personal de tareas realizadas.
- '/funcionario?tab=iniciar': Iniciar un nuevo trámite.
- '/tracking': Seguimiento de trámites.

BOTONES DE ACCIÓN DISPONIBLES (Para invocar modales, usa EXACTAMENTE la ruta indicada en el JSON):
- Pestaña 'Usuarios': Botón '+ Nuevo Usuario' (ruta: '/admin?tab=usuarios&action=new')
- Pestaña 'Clientes': Botón '+ Nuevo Cliente' (ruta: '/admin?tab=clientes&action=new')
- Pestaña 'Empresa': Botón 'Editar Perfil Institucional' (ruta: '/admin?tab=tenants&action=edit')
- Pestaña 'Departamentos': Botón '+ Nuevo Departamento' (ruta: '/admin?tab=departamentos&action=new')
- Pestaña 'Cargos': Botón '+ Nuevo Cargo' (ruta: '/admin?tab=cargos&action=new')
- Pestaña 'Formularios': Botón '+ Crear Formulario' (ruta: '/admin?tab=formularios')
- Designer Hub: Botón '+ Nuevo Proyecto' (ruta: '/designer')
- Funcionario: Botones 'Tomar Tarea', 'Completar Tarea', 'Iniciar Nuevo Proceso' (ruta: '/funcionario')

INSTRUCCIÓN CRÍTICA PARA EL FINAL DE TU RESPUESTA:
Después de tu respuesta textual, DEBES incluir EN LA ÚLTIMA LÍNEA un bloque JSON oculto con este formato exacto:
<!--NAV:{"rutaNavegacion": "/ruta/aqui", "acciones": [{"label": "Texto botón", "ruta": "/ruta"}]}-->
Si no hay navegación o no tiene acceso, usa: <!--NAV:{"rutaNavegacion": null, "acciones": []}-->
SIEMPRE incluye este bloque al final, es obligatorio."""

ASSISTANT_SYSTEM_PROMPT = """Eres 'Antigravity AI', un arquitecto de procesos BPM avanzado. Tu objetivo es ayudar al usuario a diseñar flujos de trabajo profesionales sin que tenga que usar las manos.

INSTRUCCIONES CRÍTICAS:
1. DEBES responder UNICAMENTE con un JSON válido. Sin preámbulos ni explicaciones fuera del JSON.
2. CIRUGÍA VS CREACIÓN: Si el usuario menciona un objeto que YA existe en el contexto (mira la lista de nodos abajo), NO uses CREAR_NODO. Usa MOVER_NODO, MODIFICAR_NODO o CAMBIAR_ESTILO.
3. Sé "creativo" (crear flujos completos) SOLO si el usuario pide algo nuevo o abstracto como "Crea un proceso de ventas". Si pide algo específico sobre lo que ya hay (ej: "mueve gato"), sé quirúrgico: solo mueve el nodo existente.
4. Usa nombres de actividades claros y orientados a la acción.
5. Para procesos nuevos, genera una LISTA de acciones en orden lógico.
6. Si la acción no es posible, responde en 'explicacion' y usa 'NOT_SUPPORTED'.

ESTRUCTURA DEL JSON:
{
  "explicacion": "Un mensaje empoderador y técnico de lo que vas a construir",
  "acciones": [
     { "tipo": "CREAR_CALLE", "params": { "nombre": "RRHH", "color": "#6366f1" } },
     { "tipo": "CREAR_NODO", "params": { "tipo": "TAREA", "nombre": "Entrevista Técnica", "calleNombre": "RRHH" } },
     { "tipo": "CONECTAR_NODOS", "params": { "origenNombre": "Inicio", "destinoNombre": "Entrevista Técnica" } },
     { "tipo": "ASIGNAR_PLANTILLA", "params": { "nombreNodo": "Entrevista Técnica", "nombrePlantilla": "Formulario Contratación" } }
  ]
}

ACCIONES SOPORTADAS:
- CREAR_CALLE (nombre, color)
- CREAR_NODO (tipo: INICIO|FIN|TAREA|DECISION|FORK|JOIN, nombre, calleNombre)
- ELIMINAR_NODO (nombre)
- CONECTAR_NODOS (origenNombre, destinoNombre)
- MODIFICAR_NODO (nombreActual, nuevoNombre)
- ELIMINAR_CALLE (nombre)
- MOVER_NODO (nombreNodo, nuevaCalleNombre)
- CAMBIAR_ESTILO (nombre, color, ancho, alto, fontSize: sm|md|lg)
- ASIGNAR_PLANTILLA (nombreNodo, nombrePlantilla)
- RENOMBRAR_CALLE (nombreActual, nuevoNombre)
- ELIMINAR_TRANSICION (origenNombre, destinoNombre)
- REORDENAR_CALLES (nombresOrdenados: array de strings)
- EDITAR_TRANSICION (origenNombre, destinoNombre, etiqueta, condicion, color, tipoLinea: solida|punteada|lineas, grosor)
- NOT_SUPPORTED (razon)"""


# ═══════════════════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════════════════

@app.get("/")
async def health_check():
    return {
        "service": "BPM Inteligente — AI Microservice",
        "status": "running",
        "version": "2.0.0",
        "groq_configured": bool(GROQ_API_KEY),
        "model": GROQ_MODEL,
    }


# ── 1. Asistente IA (Designer) — JSON Síncrono ──────────────────
@app.post("/api/ai/assistant/prompt", response_model=AiActionResponse)
async def assistant_prompt(request: AiCommandRequest):
    """
    Procesa instrucciones del diseñador BPM usando Groq.
    Retorna JSON estructurado con acciones para el canvas.
    NO usa streaming (el frontend necesita el JSON completo).
    """
    if not groq_client:
        return AiActionResponse(
            explicacion="API Key de Groq no configurada. Configura GROQ_API_KEY en el archivo .env del microservicio.",
            acciones=[AiAction(tipo="NOT_SUPPORTED", params={"razon": "GROQ_API_KEY no configurada en el microservicio Python."})],
        )

    try:
        # Construir contexto del diagrama si existe
        context_str = ""
        if request.contexto:
            context_str = f"\n\nContexto del diagrama actual: {json.dumps(request.contexto, ensure_ascii=False, default=str)}"

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": ASSISTANT_SYSTEM_PROMPT + context_str},
                {"role": "user", "content": request.instruccion},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
            max_tokens=1024,
        )

        content = response.choices[0].message.content
        parsed = json.loads(content)

        return AiActionResponse(
            explicacion=parsed.get("explicacion", "Procesado correctamente."),
            acciones=[AiAction(**a) for a in parsed.get("acciones", [])],
        )

    except json.JSONDecodeError as e:
        return AiActionResponse(
            explicacion=f"Error al parsear la respuesta del LLM: {str(e)}",
            acciones=[AiAction(tipo="NOT_SUPPORTED", params={"razon": "Respuesta del LLM no fue JSON válido."})],
        )
    except Exception as e:
        return AiActionResponse(
            explicacion=f"Error al procesar con Groq: {str(e)}",
            acciones=[AiAction(tipo="NOT_SUPPORTED", params={"razon": str(e)})],
        )


# ── 2. Chatbot — Streaming SSE ───────────────────────────────────
@app.post("/api/ai/chatbot/chat")
async def chatbot_chat(request: ChatbotRequest):
    """
    Chatbot conversacional con streaming SSE.
    Inyecta el contexto dinámico de Java como pseudo-RAG.
    Retorna Server-Sent Events con chunks de texto.
    """
    # Fallback sin streaming si no hay Groq
    if not groq_client:
        return ChatbotResponse(
            respuesta="API Key de Groq no configurada. Configura GROQ_API_KEY en el microservicio Python para activar el chatbot inteligente.",
            rutaNavegacion=None,
        )

    # Construir mensajes para Groq
    system_content = CHATBOT_SYSTEM_PROMPT
    if request.contextoDinamico:
        system_content += f"\n\n{request.contextoDinamico}"
    if request.contextoSeccion:
        system_content += f"\n\nEl usuario está actualmente en la sección: {request.contextoSeccion}"

    messages = [{"role": "system", "content": system_content}]

    # Añadir historial conversacional
    for msg in request.historial:
        messages.append({"role": msg.role, "content": msg.content})

    # Añadir mensaje actual
    messages.append({"role": "user", "content": request.mensaje})

    async def generate_stream():
        """Generador SSE que lee el stream de Groq y emite chunks."""
        full_response = ""
        try:
            stream = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=messages,
                temperature=0.3,
                max_tokens=2048,
                stream=True,
            )

            for chunk in stream:
                delta = chunk.choices[0].delta
                if delta.content:
                    text = delta.content
                    full_response += text
                    # Emitir chunk como SSE
                    sse_data = json.dumps({"text": text}, ensure_ascii=False)
                    yield f"data: {sse_data}\n\n"

            # Parsear navegación y acciones del bloque <!--NAV:...--> al final
            ruta_navegacion = None
            acciones = []
            visible_text = full_response

            if "<!--NAV:" in full_response:
                parts = full_response.split("<!--NAV:")
                visible_text = parts[0].strip()
                try:
                    nav_json_str = parts[1].rstrip("-->").rstrip("-").rstrip(">").strip()
                    if nav_json_str.endswith("-->"):
                        nav_json_str = nav_json_str[:-3]
                    nav_data = json.loads(nav_json_str)
                    ruta_navegacion = nav_data.get("rutaNavegacion")
                    acciones = nav_data.get("acciones", [])
                except (json.JSONDecodeError, IndexError):
                    pass

            # Evento final con metadata
            done_data = json.dumps({
                "done": True,
                "rutaNavegacion": ruta_navegacion,
                "acciones": acciones,
                "fullText": visible_text,
            }, ensure_ascii=False)
            yield f"data: {done_data}\n\n"

        except Exception as e:
            error_data = json.dumps({
                "done": True,
                "error": str(e),
                "rutaNavegacion": None,
                "acciones": [],
                "fullText": f"Error al procesar con Groq: {str(e)}",
            }, ensure_ascii=False)
            yield f"data: {error_data}\n\n"

    return StreamingResponse(
        generate_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ── 3. ML Analysis & Insights (Fase 3) ─────────────────────────
@app.post("/api/ai/ml/analyze-bottlenecks", response_model=AnalysisResponse)
async def analyze_bottlenecks(request: AnalysisRequest):
    """
    Endpoint legacy adaptado a Fase 3.
    Llama internamente a ml_service y transforma la salida al formato antiguo
    para mantener compatibilidad con Java.
    """
    try:
        # Ejecutar pipeline completo
        resultados = await ejecutar_pipeline_completo(groq_client, request.politicaId, request.tenantId)
        
        findings = []
        for cuello in resultados.get("cuellosBottella", []):
            findings.append(Finding(
                type="BOTTLENECK_REAL",
                nodeId=cuello["actividadId"],
                severity=cuello["severity"],
                message=f"Cuello de botella: Actividad '{cuello['actividadNombre']}' toma {cuello['promedioMinutos']} min (Desviación: {cuello['desviacionSobre']}x).",
                suggestion="Considera dividir esta tarea, asignar más personal, o automatizar pasos repetitivos."
            ))
            
        if not findings:
            findings.append(Finding(
                type="INFO", severity="INFO", nodeId="",
                message="El proceso opera dentro de parámetros normales.",
                suggestion="Continúe monitoreando."
            ))
            
        return AnalysisResponse(findings=findings)
        
    except Exception as e:
        import logging
        logging.error("Error en analyze_bottlenecks: %s", str(e))
        return AnalysisResponse(
            findings=[
                Finding(
                    type="ERROR", severity="CRITICAL", nodeId="",
                    message=f"Error en motor ML: {str(e)}",
                    suggestion="Revise los logs del microservicio."
                )
            ]
        )

@app.get("/api/ai/ml/insights", response_model=InsightsResponse)
@app.post("/api/ai/ml/insights", response_model=InsightsResponse)
async def get_insights(politicaId: Optional[str] = None, tenantId: Optional[str] = None):
    """
    Nuevo endpoint de Insights (Fase 3).
    Acepta GET (query param) o POST.
    Retorna el JSON completo con métricas, cuellos de botella y predicciones ML.
    """
    try:
        resultado = await ejecutar_pipeline_completo(groq_client, politicaId, tenantId)
        return InsightsResponse(**resultado)
    except Exception as e:
        import logging
        logging.error("Error obteniendo insights: %s", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/ai/ml/train")
async def train_tf_models():
    """
    Triggers retraining of TensorFlow models and reloads them in memory.
    """
    try:
        import sys
        import os
        
        # Agregar el directorio de scripts al path si no está
        scripts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts")
        if scripts_dir not in sys.path:
            sys.path.append(scripts_dir)
            
        from train_tf_model import extraer_datos, preparar_features, entrenar_modelos, exportar_artefactos
        
        import logging
        logging.info("🔄 Invocando re-entrenamiento de TensorFlow...")
        df_tramites, df_registros = extraer_datos()
        df_features = preparar_features(df_tramites, df_registros)
        
        model_mo, autoencoder, scaler, le_depto, le_politica, le_actividad = entrenar_modelos(df_features)
        exportar_artefactos(model_mo, autoencoder, scaler, le_depto, le_politica, le_actividad)
        
        # Recargar en caliente
        global tf_routing_model, tf_autoencoder, tf_preprocessing
        tf_routing_model = model_mo
        tf_autoencoder = autoencoder
        tf_preprocessing = {
            "scaler": scaler,
            "le_depto": le_depto,
            "le_politica": le_politica,
            "le_actividad": le_actividad
        }
        
        logging.info("✅ Modelos TensorFlow re-entrenados y recargados en caliente con éxito.")
        return {"status": "success", "message": "TensorFlow models retrained and hot-reloaded successfully."}
    except Exception as e:
        import logging
        logging.exception("Error en train_tf_models")
        raise HTTPException(status_code=500, detail=f"Error durante re-entrenamiento: {str(e)}")


from fastapi import Response

@app.get("/api/ai/tts")
async def text_to_speech(text: str, voice_id: Optional[str] = None):
    """
    Convierte texto a voz usando ElevenLabs. Si ElevenLabs falla (por restricciones de IP en Cloud Run),
    realiza un fallback automático a edge-tts para mantener calidad premium sin costo.
    """
    final_voice_id = voice_id or ELEVENLABS_VOICE_ID
    
    if el_client:
        try:
            audio_iter = el_client.generate(
                text=text,
                voice=final_voice_id,
                model="eleven_multilingual_v2"
            )
            audio_data = b"".join(audio_iter)
            return Response(content=audio_data, media_type="audio/mpeg")
        except Exception as e:
            import logging
            logging.error("Error en TTS ElevenLabs (fallback a edge-tts): %s", str(e))
            # No levantar error, pasar al fallback
            
    # --- FALLBACK A EDGE-TTS (Voz neuronal de Microsoft, alta calidad y gratis) ---
    try:
        import edge_tts
        import io
        # 'es-ES-ElviraNeural' o 'es-MX-DaliaNeural' son excelentes opciones
        voice = "es-ES-ElviraNeural" 
        communicate = edge_tts.Communicate(text, voice)
        
        audio_data = bytearray()
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_data.extend(chunk["data"])
                
        return Response(content=bytes(audio_data), media_type="audio/mpeg")
    except Exception as e:
        import logging
        logging.error("Error en TTS Edge: %s", str(e))
        raise HTTPException(status_code=500, detail=f"Error generando audio: {str(e)}")


@app.post("/api/ai/stt")
async def speech_to_text(file: UploadFile = File(...)):
    """
    Recibe un archivo de audio (ej. webm) y lo transcribe usando Whisper de Groq.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
    
    try:
        # Groq requiere nombre de archivo con extensión reconocida
        file_name = file.filename or "audio.webm"
        file_content = await file.read()
        
        # Llamar a Whisper API
        transcription = groq_client.audio.transcriptions.create(
            file=(file_name, file_content),
            model="whisper-large-v3",
            prompt="Transcribe el siguiente audio en español.",
            response_format="text",
            language="es"
        )
        return {"text": transcription}
    except Exception as e:
        import logging
        logging.error("Error en STT Whisper: %s", str(e))
        raise HTTPException(status_code=500, detail=f"Error al procesar audio: {str(e)}")


# ── 4. Enrutador Móvil (Voz a Acción) ─────────────────────────
@app.post("/api/ai/mobile/voice-router", response_model=VoiceRouterResponse)
async def mobile_voice_router(file: UploadFile = File(...)):
    """
    Recibe un audio del celular, transcribe con Whisper (Groq)
    y usa Llama 3 para inferir el 'politicaId' más adecuado del catálogo.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
    
    try:
        # 1. Transcripción ultrarrápida
        file_name = file.filename or "voice_query.m4a"
        file_content = await file.read()
        
        transcription = groq_client.audio.transcriptions.create(
            file=(file_name, file_content),
            model="whisper-large-v3",
            prompt="Transcribe el siguiente audio de un cliente de la cooperativa en español latino.",
            response_format="text",
            language="es"
        )
        
        # 2. Catálogo de Políticas en caché (Mock dinámico)
        # En producción esto consulta a Redis o al core de Spring Boot
        catalogo_politicas = [
            {"id": "POL-100", "nombre": "Mantenimiento Urgente", "desc": "Cortocircuitos, chispas, postes caídos o sin luz."},
            {"id": "POL-200", "nombre": "Instalación de Nuevo Medidor", "desc": "Clientes que desean conexión nueva."},
            {"id": "POL-300", "nombre": "Reclamo de Facturación", "desc": "Facturas elevadas, lectura errónea o dudas de cobro."}
        ]
        
        # 3. Razonamiento Semántico
        system_prompt = f"""Eres el cerebro enrutador del sistema BPM Inteligente.
Analiza la siguiente transcripción de un cliente e infiere la intención.
Compara la intención contra este catálogo de trámites disponibles:
{json.dumps(catalogo_politicas, ensure_ascii=False)}

RESPONDE ÚNICAMENTE CON UN JSON EN ESTE FORMATO ESTRICTO:
{{"politicaId": "AQUÍ_EL_ID", "intencionDetectada": "Breve descripción", "confianza": 0.99}}
Si no tiene sentido con ninguna, usa "politicaId": "UNKNOWN".
"""

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": transcription},
            ],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        
        return VoiceRouterResponse(
            politicaId=result.get("politicaId", "UNKNOWN"),
            intencionDetectada=result.get("intencionDetectada", "Trámite General"),
            confianza=float(result.get("confianza", 0.5))
        )
        
    except Exception as e:
        import logging
        logging.error("Error NLP Voice Router: %s", str(e))
        raise HTTPException(status_code=500, detail=f"Error analizando voz: {str(e)}")


# ── 5. Auto-llenado de Formularios por Voz ─────────────────────────
@app.post("/api/ai/forms/voice-fill")
async def voice_fill_form(
    file: UploadFile = File(...),
    fields: str = Form(...)  # JSON string of fields metadata
):
    """
    Recibe un audio de dictado de voz, lo transcribe usando Whisper,
    y usa Groq NLP para extraer estructuradamente la información de los campos del formulario.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
    
    try:
        # 1. Transcribir el audio
        file_name = file.filename or "form_voice.webm"
        file_content = await file.read()
        
        transcription = groq_client.audio.transcriptions.create(
            file=(file_name, file_content),
            model="whisper-large-v3",
            prompt="Transcribe el siguiente audio en español donde un usuario describe datos de un formulario de inspección o trámite.",
            response_format="text",
            language="es",
            temperature=0.0
        )
        
        # 2. Parsear el esquema de campos
        try:
            fields_schema = json.loads(fields)
        except Exception:
            raise HTTPException(status_code=400, detail="El campo 'fields' debe ser un JSON válido.")
        
        # 3. Prompt de extracción estructurada
        system_prompt = f"""Eres un motor de extracción de datos por inteligencia artificial de alta precisión.
Tu trabajo es procesar la siguiente transcripción de voz de un usuario y rellenar un formulario estructurado.

Esquema del formulario (Lista de campos a rellenar):
{json.dumps(fields_schema, ensure_ascii=False, indent=2)}

Reglas de extracción:
1. Extrae únicamente los valores solicitados en el esquema.
2. Cada campo en el esquema tiene un 'name' (nombre de la variable) y un 'type' (tipo de dato, ej: number, string, boolean).
3. Asegúrate de convertir los valores al tipo de dato correcto. Si es boolean, responde true/false. Si es number, responde con el número correspondiente.
4. Si un valor no se menciona o no se puede inferir con certeza, usa null.
5. Devuelve ÚNICAMENTE un objeto JSON plano donde las claves sean los 'name' del esquema y los valores sean los datos extraídos.
6. NO incluyas ninguna explicación, texto adicional, ni bloques de código de markdown. Solo el JSON plano.
7. CRÍTICO: Si un campo tiene una lista de 'options', el valor extraído DEBE ser EXACTAMENTE una de esas opciones. Si el usuario usa un sinónimo (ej: dice "positiva" y las opciones son ["Aprobado", "Rechazado"]), infiere y mapea a la opción correcta ("Aprobado"). NUNCA devuelvas un valor que no esté en la lista de options si esta existe.
"""

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Transcripción del usuario: \"{transcription}\""},
            ],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        
        extracted_values = json.loads(response.choices[0].message.content)
        
        return {
            "transcription": transcription,
            "values": extracted_values
        }
        
    except Exception as e:
        import logging
        logging.error("Error en voice-fill: %s", str(e))
        raise HTTPException(status_code=500, detail=f"Error en auto-llenado por voz: {str(e)}")
# ── 5.5 Asignación Inteligente de Trámites por IA ──────────────────────
class AiAssignRequest(BaseModel):
    prompt: str
    available_workflows: list

@app.post("/api/ai/forms/assign-workflow")
async def ai_assign_workflow(req: AiAssignRequest):
    """
    Analiza un prompt del usuario y selecciona la política de negocio más adecuada.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
    
    system_prompt = f"""Eres un asistente de Inteligencia Artificial para una empresa.
Tu trabajo es escuchar/leer el problema o solicitud de un cliente y asignarle el trámite correcto.

Lista de trámites (flujos de trabajo) disponibles en la empresa:
{json.dumps(req.available_workflows, ensure_ascii=False, indent=2)}

Reglas críticas:
1. Analiza el 'prompt' del cliente.
2. Compara su necesidad con las descripciones y nombres de los trámites disponibles.
3. Si existe un trámite que resuelve su problema, devuelve su 'id' exacto.
4. Si el cliente pide algo que la empresa NO hace (ej. pedir una pizza a una compañía eléctrica), debes rechazarlo. No asignes trámites al azar.
5. Devuelve ÚNICAMENTE un JSON con esta estructura:
   {{
     "assigned_politica_id": "string o null",
     "reason": "Un mensaje breve y amable explicando tu decisión al cliente."
   }}
6. Si decides que ningún trámite aplica, assigned_politica_id debe ser null.
"""

    try:
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Solicitud del cliente: \"{req.prompt}\""},
            ],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        import logging
        logging.error("Error en assign-workflow: %s", str(e))
        raise HTTPException(status_code=500, detail="Error en el análisis de asignación de IA.")


# ── 6. Generador de Reportes NLP Dinámicos ─────────────────────────
class NlToAggregationRequest(BaseModel):
    query: str
    politicaId: Optional[str] = None
    tenantId: Optional[str] = None

@app.post("/api/ai/ml/nl-to-aggregation")
async def nl_to_aggregation(req: NlToAggregationRequest):
    """
    Traduce una consulta en lenguaje natural a un pipeline de agregación de MongoDB seguro,
    lo ejecuta directamente sobre la base de datos y retorna los resultados formateados.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
    
    from database import get_db
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="No se pudo conectar a la base de datos.")
        
    try:
        # 1. Definir el esquema y contexto de colecciones
        schemas_info = """
Colecciones disponibles en MongoDB:

1. 'registros_actividad': Registro histórico detallado de la ejecución de cada tarea/actividad.
   Estructura:
   - tramiteId: string (ID del trámite al que pertenece)
   - actividadId: string (ID de la actividad)
   - actividadNombre: string (Nombre legible de la actividad)
   - departamentoId: string (ID del departamento asignado)
   - tenantId: string (ID del tenant/empresa)
   - estado: string ("HECHO", "EN_PROGRESO", "CANCELADO")
   - ejecutadoPor: string (Nombre del usuario que la completó)
   - ejecutadoPorId: string (ID del usuario)
   - asignadoEn: Date/ISODate (Fecha de inicio/asignación de la actividad)
   - completadoEn: Date/ISODate (Fecha de finalización de la actividad)

2. 'tramites': Instancias de procesos de trámites iniciados por clientes.
   Estructura:
   - _id: string (ID único del trámite)
   - politicaId: string (ID del flujo/política de negocio)
   - tenantId: string (ID del tenant/empresa)
   - estado: string ("COMPLETADO", "EN_PROGRESO", "CANCELADO")
   - iniciadoEn: Date/ISODate (Fecha de inicio del trámite)
   - finalizadoEn: Date/ISODate (Fecha de finalización)
   - codigoSeguimiento: string
   - documentoCliente: string (CI/DNI del cliente)
   - clienteNombre: string (Nombre del cliente)

3. 'clientes': Base de datos de clientes externos.
   Estructura:
   - _id: string (CI o ID único del cliente)
   - tenantId: string (ID del tenant/empresa)
   - nombre: string
   - apellido: string
   - ci: string
   - correo: string
   - telefono: string
   - direccion: string
   - creadoEn: Date/ISODate

4. 'departamentos': Departamentos de la organización.
   Estructura:
   - _id: string
   - tenantId: string
   - nombre: string
   - codigo: string
   - ubicacion: string
   - presupuesto: number
"""

        system_prompt = f"""Eres un Ingeniero de Datos y experto en MongoDB de alta precisión.
Tu tarea es traducir una consulta en lenguaje natural del usuario a una consulta agregada de MongoDB (Aggregation Pipeline).

Contexto del esquema de la Base de Datos:
{schemas_info}

Reglas críticas de generación:
1. Analiza el requerimiento del usuario y selecciona la colección principal más adecuada (ej: 'registros_actividad', 'tramites', 'clientes').
2. Diseña un pipeline de agregación válido como un array de etapas MongoDB (ej: [{{"$match": ...}}, {{"$group": ...}}, {{"$sort": ...}}]).
3. Usa operadores estándar de agregación de MongoDB ($match, $group, $sort, $limit, $project, $lookup, $unwind).
4. No uses ninguna operación de escritura ($out, $merge, $write, $destroy).
5. Responde estrictamente con un objeto JSON con este formato exacto:
{{
  "collection": "nombre_de_la_coleccion",
  "pipeline": [
    {{ "$match": ... }},
    {{ "$group": ... }}
  ]
}}
6. NO agregues explicaciones ni código markdown. Solo devuelve el objeto JSON válido.
"""

        # Agregar contexto de consulta
        user_content = f"Consulta: \"{req.query}\"\n"
        if req.politicaId:
            user_content += f"Filtrar por politicaId: \"{req.politicaId}\"\n"
        if req.tenantId:
            user_content += f"Filtrar por tenantId: \"{req.tenantId}\"\n"

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        
        result_json = json.loads(response.choices[0].message.content)
        collection_name = result_json.get("collection")
        pipeline = result_json.get("pipeline", [])
        
        # ── 4. Validación de Seguridad Estricta (Sandbox) ───────────────
        allowed_collections = ["tramites", "registros_actividad", "clientes", "departamentos", "usuarios", "audit_log", "politicas_negocio"]
        if collection_name not in allowed_collections:
            raise HTTPException(
                status_code=400, 
                detail=f"Colección no permitida: {collection_name}. Las permitidas son: {allowed_collections}"
            )
            
        # Buscar operadores destructivos serializando el pipeline a texto
        pipeline_str = json.dumps(pipeline).lower()
        destructive_ops = ["$out", "$merge", "$write", "$eval", "$runcommand", "$accumulator", "$function"]
        for op in destructive_ops:
            if op in pipeline_str:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Operador no permitido detectado en la agregación: {op}"
                )
                
        # Asegurar el filtro de tenantId para seguridad y aislamiento
        if req.tenantId:
            # Buscar si ya hay un $match inicial, si no, crearlo.
            tenant_match = {"tenantId": req.tenantId}
            # Si req.politicaId también está presente, agregarlo.
            if req.politicaId and collection_name in ["tramites", "registros_actividad"]:
                tenant_match["politicaId"] = req.politicaId
                
            has_match = False
            for stage in pipeline:
                if "$match" in stage:
                    stage["$match"].update(tenant_match)
                    has_match = True
                    break
            if not has_match:
                pipeline.insert(0, {"$match": tenant_match})
                
        # 5. Ejecutar la agregación en MongoDB
        import asyncio
        coll = db[collection_name]
        
        # Ejecutar en thread pool para no bloquear el bucle de eventos asíncrono
        def execute_aggregation():
            return list(coll.aggregate(pipeline))
            
        raw_results = await asyncio.to_thread(execute_aggregation)
        
        # Helper para serializar tipos especiales de BSON a JSON estándar (como ObjectId y datetime)
        def json_serializable(data):
            if isinstance(data, list):
                return [json_serializable(item) for item in data]
            elif isinstance(data, dict):
                return {k: json_serializable(v) for k, v in data.items()}
            elif hasattr(data, "isoformat"):
                return data.isoformat()
            elif hasattr(data, "__str__") and data.__class__.__name__ == "ObjectId":
                return str(data)
            else:
                return data
                
        serialized_results = json_serializable(raw_results)
        
        return {
            "collection": collection_name,
            "pipeline": pipeline,
            "results": serialized_results
        }
        
    except HTTPException:
        raise
    except Exception as e:
        import logging
        logging.exception("Error en nl-to-aggregation")
        raise HTTPException(status_code=500, detail=f"Error en el generador de reportes NLP: {str(e)}")



# ── Reportes Dinámicos Exportables (Fase 4 & Adiciones) ───────────

class ExportReportRequest(BaseModel):
    query: Optional[str] = None
    collection: Optional[str] = None
    pipeline: Optional[List[dict]] = None
    tenantId: Optional[str] = None
    politicaId: Optional[str] = None
    format: str  # "pdf", "xlsx", "docx"
    chartImage: Optional[str] = None

def DateNowStr() -> str:
    from datetime import datetime
    return datetime.utcnow().strftime("%Y%m%d%H%M%S")

def generate_excel_report(data: list, title: str) -> io.BytesIO:
    import io
    import pandas as pd
    df = pd.DataFrame(data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Resultados')
        workbook = writer.book
        worksheet = writer.sheets['Resultados']
        # Auto-ajustar columnas
        for col in worksheet.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = col[0].column_letter
            worksheet.column_dimensions[col_letter].width = max(max_len + 3, 10)
    output.seek(0)
    return output

def generate_docx_report(data: list, title: str, query: str, tenant_id: str, politica_id: str, chart_image: str = None) -> io.BytesIO:
    import io
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Título elegante
    t = doc.add_paragraph()
    run = t.add_run("BPM INTELIGENTE - REPORTE DE ANALÍTICA NLP")
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229) # Indigo
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generado el: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    if query:
        doc.add_paragraph(f"Consulta NLP: {query}")
    if politica_id:
        doc.add_paragraph(f"Filtrado por Política ID: {politica_id}")

    if chart_image and chart_image.startswith("data:image"):
        try:
            import base64
            img_data = base64.b64decode(chart_image.split(",")[1])
            img_io = io.BytesIO(img_data)
            doc.add_heading("Visualización de Datos", level=2)
            doc.add_picture(img_io, width=Pt(400))
        except Exception as e:
            print("Error rendering chartImage for DOCX:", e)
        
    doc.add_heading("Datos Generados", level=2)
    
    if not data:
        doc.add_paragraph("No se encontraron registros.")
    else:
        keys = list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(keys))
        table.style = 'Light Shading Accent 1'
        
        # Cabecera
        hdr_cells = table.rows[0].cells
        for idx, key in enumerate(keys):
            header_text = "Categoría / Agrupación" if key == "_id" else str(key).capitalize()
            hdr_cells[idx].text = header_text
            hdr_cells[idx].paragraphs[0].runs[0].bold = True
            
        # Filas
        for row in data:
            row_cells = table.add_row().cells
            for idx, key in enumerate(keys):
                val = row.get(key, '')
                if val is None:
                    row_cells[idx].text = '-'
                elif isinstance(val, dict) or isinstance(val, list):
                    row_cells[idx].text = json.dumps(val, ensure_ascii=False)
                else:
                    row_cells[idx].text = str(val)
                    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_pdf_report(data: list, title: str, query: str, tenant_id: str, politica_id: str, chart_image: str = None) -> io.BytesIO:
    import io
    from datetime import datetime
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Estilos CSS premium en PDF
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=colors.HexColor('#4f46e5'),
        alignment=1, # Centro
        spaceAfter=15
    )
    
    meta_style = ParagraphStyle(
        'DocMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        textColor=colors.HexColor('#475569'),
        spaceAfter=5
    )
    
    header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8,
        textColor=colors.white
    )
    
    cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7,
        textColor=colors.HexColor('#0f172a')
    )
    
    story.append(Paragraph("BPM INTELIGENTE - REPORTE DE ANALÍTICA NLP", title_style))
    story.append(Paragraph(f"<b>Generado el:</b> {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}", meta_style))
    if query:
        story.append(Paragraph(f"<b>Consulta NLP:</b> {query}", meta_style))
    if politica_id:
        story.append(Paragraph(f"<b>Política de Negocio:</b> {politica_id}", meta_style))
    story.append(Spacer(1, 15))

    if chart_image and chart_image.startswith("data:image"):
        try:
            import base64
            from reportlab.platypus import Image as RLImage
            img_data = base64.b64decode(chart_image.split(",")[1])
            img_io = io.BytesIO(img_data)
            img = RLImage(img_io, width=400, height=250)
            img.hAlign = 'CENTER'
            story.append(img)
            story.append(Spacer(1, 15))
        except Exception as e:
            print("Error parsing chartImage for PDF:", e)
    
    if not data:
        story.append(Paragraph("No se encontraron registros.", cell_style))
    else:
        keys = list(data[0].keys())
        table_data = []
        
        # Fila cabecera
        hdr_row = []
        for k in keys:
            header_text = "Categoría / Agrupación" if k == "_id" else str(k).capitalize()
            hdr_row.append(Paragraph(header_text, header_style))
        table_data.append(hdr_row)
        
        # Filas de datos
        for row in data:
            row_row = []
            for k in keys:
                val = row.get(k, '')
                if val is None:
                    txt = '-'
                elif isinstance(val, dict) or isinstance(val, list):
                    txt = json.dumps(val, ensure_ascii=False)
                else:
                    txt = str(val)
                row_row.append(Paragraph(txt, cell_style))
            table_data.append(row_row)
            
        t = Table(table_data, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4f46e5')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#4f46e5')),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')])
        ]))
        story.append(t)
        
    doc.build(story)
    output.seek(0)
    return output

@app.post("/api/ai/ml/export-report")
async def export_report(req: ExportReportRequest):
    """
    Ejecuta una consulta agregada NLP y retorna los resultados formateados en un stream
    descargable de PDF, Word o Excel. Valida que existan datos reales para evitar alucinaciones.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
        
    from database import get_db
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="No se pudo conectar a la base de datos.")
        
    try:
        collection_name = req.collection
        pipeline = req.pipeline
        
        # 1. Generar pipeline si no viene especificado pero hay consulta NLP
        if not pipeline and req.query:
            schemas_info = """
Colecciones disponibles en MongoDB:

1. 'registros_actividad': Registro histórico detallado de la ejecución de cada tarea/actividad.
   Estructura:
   - tramiteId: string
   - actividadId: string
   - actividadNombre: string
   - departamentoId: string
   - tenantId: string
   - estado: string ("HECHO", "EN_PROGRESO", "CANCELADO")
   - ejecutadoPor: string
   - completadoEn: Date/ISODate

2. 'tramites': Instancias de procesos de trámites iniciados por clientes.
   Estructura:
   - _id: string
   - politicaId: string
   - tenantId: string
   - estado: string ("COMPLETADO", "EN_PROGRESO", "CANCELADO")
   - iniciadoEn: Date/ISODate
   - finalizadoEn: Date/ISODate
   - clienteNombre: string

3. 'clientes': Base de datos de clientes externos.
   Estructura:
   - _id: string
   - tenantId: string
   - nombre: string
   - apellido: string
   - correo: string

4. 'departamentos': Departamentos de la organización.
   Estructura:
   - _id: string
   - tenantId: string
   - nombre: string
   - presupuesto: number
"""
            system_prompt = f"""Eres un Ingeniero de Datos y experto en MongoDB de alta precisión.
Tu tarea es traducir una consulta en lenguaje natural del usuario a una consulta agregada de MongoDB.

Contexto del esquema de la Base de Datos:
{schemas_info}

Reglas críticas de generación:
1. Analiza el requerimiento del usuario y selecciona la colección principal más adecuada.
2. Diseña un pipeline de agregación válido como un array de etapas MongoDB.
3. Responde estrictamente con un objeto JSON con este formato exacto:
{{
  "collection": "nombre_de_la_coleccion",
  "pipeline": [
    {{ "$match": ... }}
  ]
}}
4. NO agregues explicaciones ni código markdown. Solo devuelve el objeto JSON válido.
"""
            user_content = f"Consulta: \"{req.query}\"\n"
            if req.politicaId:
                user_content += f"Filtrar por politicaId: \"{req.politicaId}\"\n"
            if req.tenantId:
                user_content += f"Filtrar por tenantId: \"{req.tenantId}\"\n"

            response = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content},
                ],
                temperature=0.0,
                response_format={"type": "json_object"}
            )
            result_json = json.loads(response.choices[0].message.content)
            collection_name = result_json.get("collection")
            pipeline = result_json.get("pipeline", [])

        if not collection_name or not pipeline:
            raise HTTPException(status_code=400, detail="Colección o consulta inválida.")

        # 2. Validación de Seguridad Estricta (Sandbox)
        allowed_collections = ["tramites", "registros_actividad", "clientes", "departamentos", "usuarios", "audit_log", "politicas_negocio"]
        if collection_name not in allowed_collections:
            raise HTTPException(
                status_code=400, 
                detail=f"Colección no permitida: {collection_name}. Las permitidas son: {allowed_collections}"
            )
            
        pipeline_str = json.dumps(pipeline).lower()
        destructive_ops = ["$out", "$merge", "$write", "$eval", "$runcommand", "$accumulator", "$function"]
        for op in destructive_ops:
            if op in pipeline_str:
                raise HTTPException(status_code=400, detail="Operador no permitido detectado en la agregación.")

        # Asegurar el filtro de tenantId para seguridad y aislamiento
        if req.tenantId:
            tenant_match = {"tenantId": req.tenantId}
            if req.politicaId and collection_name in ["tramites", "registros_actividad"]:
                tenant_match["politicaId"] = req.politicaId
                
            has_match = False
            for stage in pipeline:
                if "$match" in stage:
                    stage["$match"].update(tenant_match)
                    has_match = True
                    break
            if not has_match:
                pipeline.insert(0, {"$match": tenant_match})

        # 3. Ejecutar agregación en MongoDB
        coll = db[collection_name]
        
        def execute_aggregation():
            return list(coll.aggregate(pipeline))
            
        raw_results = await asyncio.to_thread(execute_aggregation)
        
        # ── 4. Validación para Evitar Alucinaciones ──
        if not raw_results:
            raise HTTPException(
                status_code=400, 
                detail="No se encontraron registros de datos para los filtros y criterios seleccionados. Generación de reporte cancelada para evitar alucinaciones."
            )

        # Helper para serializar tipos especiales de BSON
        def json_serializable(data):
            if isinstance(data, list):
                return [json_serializable(item) for item in data]
            elif isinstance(data, dict):
                return {k: json_serializable(v) for k, v in data.items()}
            elif hasattr(data, "isoformat"):
                return data.isoformat()
            elif hasattr(data, "__str__") and data.__class__.__name__ == "ObjectId":
                return str(data)
            else:
                return data
                
        serialized_results = json_serializable(raw_results)

        # ── Limpiar claves internas y técnicas para el usuario final (jefes/empresarios) ──
        cleaned_results = []
        for row in serialized_results:
            if not isinstance(row, dict):
                cleaned_results.append(row)
                continue
            cleaned_row = {}
            for k, v in row.items():
                if k in ["_class", "tenantId", "ejecutadoPorId", "proyectoId", "politicaId"]:
                    continue
                if k == "_id":
                    # Ocultar si es un ObjectId estándar de 24 caracteres hexadecimales
                    if isinstance(v, str) and len(v) == 24 and all(c in "0123456789abcdefABCDEF" for c in v):
                        continue
                    # Ocultar si es un objeto ObjectId serializado
                    if isinstance(v, dict) and "$oid" in v:
                        continue
                cleaned_row[k] = v
            cleaned_results.append(cleaned_row)
        serialized_results = cleaned_results

        # 5. Generar archivo y responder
        from fastapi.responses import StreamingResponse
        import io
        
        fmt = req.format.lower()
        if fmt == "xlsx":
            file_stream = generate_excel_report(serialized_results, f"Reporte de {collection_name}")
            filename = f"reporte-{collection_name}-{DateNowStr()}.xlsx"
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif fmt == "docx":
            file_stream = generate_docx_report(serialized_results, f"Reporte de {collection_name}", req.query, req.tenantId, req.politicaId, req.chartImage)
            filename = f"reporte-{collection_name}-{DateNowStr()}.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fmt == "pdf":
            file_stream = generate_pdf_report(serialized_results, f"Reporte de {collection_name}", req.query, req.tenantId, req.politicaId, req.chartImage)
            filename = f"reporte-{collection_name}-{DateNowStr()}.pdf"
            media_type = "application/pdf"
        else:
            raise HTTPException(status_code=400, detail=f"Formato no soportado: {req.format}")

        return StreamingResponse(
            file_stream,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        import logging
        logging.exception("Error en export-report")
        raise HTTPException(status_code=500, detail=f"Error al generar reporte: {str(e)}")



# ── TensorFlow Models (Fase 5) ───────────────────────────────────
tf_routing_model = None
tf_autoencoder = None
tf_preprocessing = None

@app.on_event("startup")
async def load_tf_models():
    global tf_routing_model, tf_autoencoder, tf_preprocessing
    import logging
    try:
        import tensorflow as tf
        import pickle
        import os
        
        models_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models")
        routing_path = os.path.join(models_dir, "routing_multi_output.keras")
        anomaly_path = os.path.join(models_dir, "anomaly_autoencoder.keras")
        prep_path = os.path.join(models_dir, "preprocessing_pipeline.pkl")
        
        if os.path.exists(routing_path) and os.path.exists(anomaly_path) and os.path.exists(prep_path):
            tf_routing_model = tf.keras.models.load_model(routing_path)
            tf_autoencoder = tf.keras.models.load_model(anomaly_path)
            with open(prep_path, "rb") as f:
                tf_preprocessing = pickle.load(f)
            logging.info("✅ TensorFlow Models Loaded Successfully.")
        else:
            logging.warning("⚠️ TensorFlow Models not found. Run train_tf_model.py to generate them.")
    except Exception as e:
        logging.error(f"❌ Error loading TensorFlow models: {e}")


class PredictRouteRequest(BaseModel):
    hora_del_dia: int
    dia_de_semana: int
    departamento_id: str
    politica_id: str
    carga_actual: int
    historial_cliente: float


class PredictRouteResponse(BaseModel):
    rutaSugerida: str
    tiempoEstimadoMinutos: float
    prioridadRecomendada: str
    isAnomalo: bool
    scoreEficiencia: float


@app.post("/api/ai/ml/predict-route", response_model=PredictRouteResponse)
async def predict_route(req: PredictRouteRequest):
    if tf_routing_model is None or tf_autoencoder is None or tf_preprocessing is None:
        raise HTTPException(status_code=503, detail="TensorFlow models are not loaded.")
        
    try:
        import numpy as np
        import math
        
        def safe_float(val, default=0.0):
            """Convert to float, replacing NaN/Inf with default."""
            f = float(val)
            if math.isnan(f) or math.isinf(f):
                return default
            return f
        
        # 1. Extraer preprocesadores
        scaler = tf_preprocessing["scaler"]
        le_depto = tf_preprocessing["le_depto"]
        le_politica = tf_preprocessing["le_politica"]
        le_actividad = tf_preprocessing["le_actividad"]
        
        # 2. Transformar entradas seguras
        try:
            depto_idx = le_depto.transform([req.departamento_id])[0]
        except:
            depto_idx = 0
            
        try:
            pol_idx = le_politica.transform([req.politica_id])[0]
        except:
            pol_idx = 0
            
        # 3. Vectorización
        features = np.array([[
            req.hora_del_dia,
            req.dia_de_semana,
            depto_idx,
            pol_idx,
            req.carga_actual,
            req.historial_cliente
        ]])
        
        features_scaled = scaler.transform(features)
        
        # Replace any NaN in scaled features with 0
        features_scaled = np.nan_to_num(features_scaled, nan=0.0, posinf=0.0, neginf=0.0)
        
        # 4. Inferencia Multi-Output
        preds = tf_routing_model.predict(features_scaled, verbose=0)
        
        duracion_pred = safe_float(preds[0][0][0], 30.0)
        if duracion_pred < 0: duracion_pred = 0.0
            
        prioridad_idx = int(np.argmax(preds[1][0]))
        map_prioridad = {0: "BAJA", 1: "MEDIA", 2: "ALTA"}
        prioridad_str = map_prioridad.get(prioridad_idx, "MEDIA")
        
        ruta_idx = int(np.argmax(preds[2][0]))
        try:
            ruta_str = le_actividad.inverse_transform([ruta_idx])[0]
        except:
            ruta_str = "DEFAULT_NEXT"
            
        # 5. Detección de Anomalías (Autoencoder)
        reconstruction = tf_autoencoder.predict(features_scaled, verbose=0)
        mse = np.mean(np.power(features_scaled - reconstruction, 2), axis=1)
        mse_val = safe_float(mse[0], 0.0)
        is_anomalo = bool(mse_val > 2.5)
        
        # 6. Calcular Score de Eficiencia (Heurística)
        score = 1.0 - (duracion_pred / 300.0)
        score = max(0.1, min(0.99, score))
        
        return PredictRouteResponse(
            rutaSugerida=ruta_str,
            tiempoEstimadoMinutos=round(duracion_pred, 1),
            prioridadRecomendada=prioridad_str,
            isAnomalo=is_anomalo,
            scoreEficiencia=round(score, 2)
        )
        
    except Exception as e:
        import logging
        logging.error(f"Error in predict_route: {str(e)}")
        raise HTTPException(status_code=500, detail="Error durante inferencia de TensorFlow.")

# ═══════════════════════════════════════════════════════════════════
# PUNTO DE ENTRADA
# ═══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", os.getenv("AI_SERVICE_PORT", 8000)))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=False)
