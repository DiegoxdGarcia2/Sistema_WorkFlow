import json
import asyncio
import io
from fastapi import APIRouter, HTTPException, Response
from fastapi.responses import StreamingResponse
from typing import Optional, List
from app.core.config import groq_client, GROQ_MODEL
import app.core.config as config
from app.schemas import (
    AnalysisRequest,
    AnalysisResponse,
    InsightsResponse,
    Finding,
    NlToAggregationRequest,
    ExportReportRequest,
    PredictRouteRequest,
    PredictRouteResponse,
)
from ml_service import ejecutar_pipeline_completo

router = APIRouter()

# Helper para reporte dinámico (Date String)
def DateNowStr() -> str:
    from datetime import datetime
    return datetime.utcnow().strftime("%Y%m%d%H%M%S")

# Excel Generator helper
def generate_excel_report(data: list, title: str) -> io.BytesIO:
    import pandas as pd
    df = pd.DataFrame(data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Resultados')
        workbook = writer.book
        worksheet = writer.sheets['Resultados']
        # Auto-ajustar columnas
        for col in worksheet.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = col[0].column_letter
            worksheet.column_dimensions[col_letter].width = max(max_len + 3, 10)
    output.seek(0)
    return output

# Docx Generator helper
def generate_docx_report(data: list, title: str, query: str, tenant_id: str, politica_id: str, chart_image: str = None) -> io.BytesIO:
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    t = doc.add_paragraph()
    run = t.add_run("BPM INTELIGENTE - REPORTE DE ANALÍTICA NLP")
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229) # Indigo
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generado el: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    if query:
        doc.add_paragraph(f"Consulta NLP: {query}")
    if politica_id:
        doc.add_paragraph(f"Filtrado por Política ID: {politica_id}")

    if chart_image and chart_image.startswith("data:image"):
        try:
            import base64
            img_data = base64.b64decode(chart_image.split(",")[1])
            img_io = io.BytesIO(img_data)
            doc.add_heading("Visualización de Datos", level=2)
            doc.add_picture(img_io, width=Pt(400))
        except Exception as e:
            print("Error rendering chartImage for DOCX:", e)
        
    doc.add_heading("Datos Generados", level=2)
    
    if not data:
        doc.add_paragraph("No se encontraron registros.")
    else:
        keys = list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(keys))
        table.style = 'Light Shading Accent 1'
        
        hdr_cells = table.rows[0].cells
        for idx, key in enumerate(keys):
            header_text = "Categoría / Agrupación" if key == "_id" else str(key).capitalize()
            hdr_cells[idx].text = header_text
            hdr_cells[idx].paragraphs[0].runs[0].bold = True
            
        for row in data:
            row_cells = table.add_row().cells
            for idx, key in enumerate(keys):
                val = row.get(key, '')
                if val is None:
                    row_cells[idx].text = '-'
                elif isinstance(val, dict) or isinstance(val, list):
                    row_cells[idx].text = json.dumps(val, ensure_ascii=False)
                else:
                    row_cells[idx].text = str(val)
                    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# PDF Generator helper
def generate_pdf_report(data: list, title: str, query: str, tenant_id: str, politica_id: str, chart_image: str = None) -> io.BytesIO:
    from datetime import datetime
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = []
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=colors.HexColor('#4f46e5'),
        alignment=1,
        spaceAfter=15
    )
    
    meta_style = ParagraphStyle(
        'DocMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        textColor=colors.HexColor('#475569'),
        spaceAfter=5
    )
    
    header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8,
        textColor=colors.white
    )
    
    cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7,
        textColor=colors.HexColor('#0f172a')
    )
    
    story.append(Paragraph("BPM INTELIGENTE - REPORTE DE ANALÍTICA NLP", title_style))
    story.append(Paragraph(f"<b>Generado el:</b> {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}", meta_style))
    if query:
        story.append(Paragraph(f"<b>Consulta NLP:</b> {query}", meta_style))
    if politica_id:
        story.append(Paragraph(f"<b>Política de Negocio:</b> {politica_id}", meta_style))
    story.append(Spacer(1, 15))

    if chart_image and chart_image.startswith("data:image"):
        try:
            import base64
            from reportlab.platypus import Image as RLImage
            img_data = base64.b64decode(chart_image.split(",")[1])
            img_io = io.BytesIO(img_data)
            img = RLImage(img_io, width=400, height=250)
            img.hAlign = 'CENTER'
            story.append(img)
            story.append(Spacer(1, 15))
        except Exception as e:
            print("Error parsing chartImage for PDF:", e)
    
    if not data:
        story.append(Paragraph("No se encontraron registros.", cell_style))
    else:
        keys = list(data[0].keys())
        table_data = []
        
        hdr_row = []
        for k in keys:
            header_text = "Categoría / Agrupación" if k == "_id" else str(k).capitalize()
            hdr_row.append(Paragraph(header_text, header_style))
        table_data.append(hdr_row)
        
        for row in data:
            row_row = []
            for k in keys:
                val = row.get(k, '')
                if val is None:
                    txt = '-'
                elif isinstance(val, dict) or isinstance(val, list):
                    txt = json.dumps(val, ensure_ascii=False)
                else:
                    txt = str(val)
                row_row.append(Paragraph(txt, cell_style))
            table_data.append(row_row)
            
        t = Table(table_data, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4f46e5')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#4f46e5')),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')])
        ]))
        story.append(t)
        
    doc.build(story)
    output.seek(0)
    return output


@router.post("/analyze-bottlenecks", response_model=AnalysisResponse)
async def analyze_bottlenecks(request: AnalysisRequest):
    """
    Endpoint legacy adaptado a Fase 3.
    Llama internamente a ml_service y transforma la salida al formato antiguo
    para mantener compatibilidad con Java.
    """
    try:
        resultados = await ejecutar_pipeline_completo(groq_client, request.politicaId, request.tenantId)
        
        findings = []
        for cuello in resultados.get("cuellosBottella", []):
            findings.append(Finding(
                type="BOTTLENECK_REAL",
                nodeId=cuello["actividadId"],
                severity=cuello["severity"],
                message=f"Cuello de botella: Actividad '{cuello['actividadNombre']}' toma {cuello['promedioMinutos']} min (Desviación: {cuello['desviacionSobre']}x).",
                suggestion="Considera dividir esta tarea, asignar más personal, o automatizar pasos repetitivos."
            ))
            
        if not findings:
            findings.append(Finding(
                type="INFO", severity="INFO", nodeId="",
                message="El proceso opera dentro de parámetros normales.",
                suggestion="Continúe monitoreando."
            ))
            
        return AnalysisResponse(findings=findings)
        
    except Exception as e:
        import logging
        logging.error("Error en analyze_bottlenecks: %s", str(e))
        return AnalysisResponse(
            findings=[
                Finding(
                    type="ERROR", severity="CRITICAL", nodeId="",
                    message=f"Error en motor ML: {str(e)}",
                    suggestion="Revise los logs del microservicio."
                )
            ]
        )

@router.get("/insights", response_model=InsightsResponse)
@router.post("/insights", response_model=InsightsResponse)
async def get_insights(politicaId: Optional[str] = None, tenantId: Optional[str] = None):
    """
    Nuevo endpoint de Insights (Fase 3).
    Acepta GET (query param) o POST.
    """
    try:
        resultado = await ejecutar_pipeline_completo(groq_client, politicaId, tenantId)
        return InsightsResponse(**resultado)
    except Exception as e:
        import logging
        logging.error("Error obteniendo insights: %s", str(e))
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/train")
async def train_tf_models():
    """
    Triggers retraining of TensorFlow models and reloads them in memory.
    """
    try:
        import sys
        import os
        
        # Agregar el directorio de scripts al path si no está
        parent_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
        scripts_dir = os.path.join(parent_dir, "scripts")
        if scripts_dir not in sys.path:
            sys.path.append(scripts_dir)
            
        from train_tf_model import extraer_datos, preparar_features, entrenar_modelos, exportar_artefactos
        
        import logging
        logging.info("🔄 Invocando re-entrenamiento de TensorFlow...")
        df_tramites, df_registros = extraer_datos()
        df_features = preparar_features(df_tramites, df_registros)
        
        model_mo, autoencoder, scaler, le_depto, le_politica, le_actividad, metrics = entrenar_modelos(df_features)
        exportar_artefactos(model_mo, autoencoder, scaler, le_depto, le_politica, le_actividad)
        
        import tensorflow as tf
        import tf2onnx
        import onnx
        import onnxruntime as ort
        
        models_dir = os.path.join(parent_dir, "models")
        routing_onnx_path = os.path.join(models_dir, "routing_multi_output.onnx")
        anomaly_onnx_path = os.path.join(models_dir, "anomaly_autoencoder.onnx")
        
        spec = (tf.TensorSpec((None, 6), tf.float32, name="input_features"),)
        routing_onnx, _ = tf2onnx.convert.from_keras(model_mo, input_signature=spec, opset=13)
        onnx.save(routing_onnx, routing_onnx_path)
        
        anomaly_onnx, _ = tf2onnx.convert.from_keras(autoencoder, input_signature=spec, opset=13)
        onnx.save(anomaly_onnx, anomaly_onnx_path)
        
        sess_opts = ort.SessionOptions()
        sess_opts.intra_op_num_threads = 1
        sess_opts.inter_op_num_threads = 1
        
        config.ort_routing_session = ort.InferenceSession(routing_onnx_path, sess_opts)
        config.ort_anomaly_session = ort.InferenceSession(anomaly_onnx_path, sess_opts)
        
        config.ort_routing_input_name = config.ort_routing_session.get_inputs()[0].name
        config.ort_anomaly_input_name = config.ort_anomaly_session.get_inputs()[0].name
        
        config.tf_preprocessing = {
            "scaler": scaler,
            "le_depto": le_depto,
            "le_politica": le_politica,
            "le_actividad": le_actividad
        }
        
        logging.info("✅ Modelos Keras re-entrenados, convertidos a ONNX y recargados en caliente con éxito.")
        return {
            "status": "success", 
            "message": "Keras models retrained, converted to ONNX, and hot-reloaded successfully.",
            "metrics": metrics
        }
    except Exception as e:
        import logging
        logging.exception("Error en train_tf_models")
        raise HTTPException(status_code=500, detail=f"Error durante re-entrenamiento: {str(e)}")

@router.post("/nl-to-aggregation")
async def nl_to_aggregation(req: NlToAggregationRequest):
    """
    Traduce una consulta en lenguaje natural a un pipeline de agregación de MongoDB seguro,
    lo ejecuta directamente sobre la base de datos y retorna los resultados formateados.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
    
    from database import get_db
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="No se pudo conectar a la base de datos.")
        
    try:
        schemas_info = """
Colecciones disponibles en MongoDB:

1. 'registros_actividad': Registro histórico detallado de la ejecución de cada tarea/actividad.
   Estructura:
   - tramiteId: string (ID del trámite al que pertenece)
   - actividadId: string (ID de la actividad)
   - actividadNombre: string (Nombre legible de la actividad)
   - departamentoId: string (ID del departamento asignado)
   - tenantId: string (ID del tenant/empresa)
   - estado: string ("HECHO", "EN_PROGRESO", "CANCELADO")
   - ejecutadoPor: string (Nombre del usuario que la completó)
   - ejecutadoPorId: string (ID del usuario)
   - asignadoEn: Date/ISODate (Fecha de inicio/asignación de la actividad)
   - completadoEn: Date/ISODate (Fecha de finalización de la actividad)

2. 'tramites': Instancias de procesos de trámites iniciados por clientes.
   Estructura:
   - _id: string (ID único del trámite)
   - politicaId: string (ID del flujo/política de negocio)
   - tenantId: string (ID del tenant/empresa)
   - estado: string ("COMPLETADO", "EN_PROGRESO", "CANCELADO")
   - iniciadoEn: Date/ISODate (Fecha de inicio del trámite)
   - finalizadoEn: Date/ISODate (Fecha de finalización)
   - codigoSeguimiento: string
   - documentoCliente: string (CI/DNI del cliente)
   - clienteNombre: string (Nombre del cliente)

3. 'clientes': Base de datos de clientes externos.
   Estructura:
   - _id: string (CI o ID único del cliente)
   - tenantId: string (ID del tenant/empresa)
   - nombre: string
   - apellido: string
   - ci: string
   - correo: string
   - telefono: string
   - direccion: string
   - creadoEn: Date/ISODate

4. 'departamentos': Departamentos de la organización.
   Estructura:
   - _id: string
   - tenantId: string
   - nombre: string
   - codigo: string
   - ubicacion: string
   - presupuesto: number
"""

        system_prompt = f"""Eres un Ingeniero de Datos y experto en MongoDB de alta precisión.
Tu tarea es traducir una consulta en lenguaje natural del usuario a una consulta agregada de MongoDB (Aggregation Pipeline).

Contexto del esquema de la Base de Datos:
{schemas_info}

Reglas críticas de generación:
1. Analiza el requerimiento del usuario y selecciona la colección principal más adecuada (ej: 'registros_actividad', 'tramites', 'clientes').
2. Diseña un pipeline de agregación válido como un array de etapas MongoDB (ej: [{{"$match": ...}}, {{"$group": ...}}, {{"$sort": ...}}]).
3. Usa operadores estándar de agregación de MongoDB ($match, $group, $sort, $limit, $project, $lookup, $unwind).
4. No uses ninguna operación de escritura ($out, $merge, $write, $destroy).
5. Responde estrictamente con un objeto JSON con este formato exacto:
{{
  "collection": "nombre_de_la_coleccion",
  "pipeline": [
    {{ "$match": ... }},
    {{ "$group": ... }}
  ]
}}
6. NO agregues explicaciones ni código markdown. Solo devuelve el objeto JSON válido.
"""

        user_content = f"Consulta: \"{req.query}\"\n"
        if req.politicaId:
            user_content += f"Filtrar por politicaId: \"{req.politicaId}\"\n"
        if req.tenantId:
            user_content += f"Filtrar por tenantId: \"{req.tenantId}\"\n"

        messages = [{"role": "system", "content": system_prompt}]
        if req.history:
            for msg in req.history:
                role = "assistant" if msg.role in ("assistant", "system") else "user"
                messages.append({"role": role, "content": msg.content})
        messages.append({"role": "user", "content": user_content})

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=messages,
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        
        result_json = json.loads(response.choices[0].message.content)
        collection_name = result_json.get("collection")
        pipeline = result_json.get("pipeline", [])
        
        allowed_collections = ["tramites", "registros_actividad", "clientes", "departamentos", "usuarios", "audit_log", "politicas_negocio"]
        if collection_name not in allowed_collections:
            raise HTTPException(
                status_code=400, 
                detail=f"Colección no permitida: {collection_name}. Las permitidas son: {allowed_collections}"
            )
            
        pipeline_str = json.dumps(pipeline).lower()
        destructive_ops = ["$out", "$merge", "$write", "$eval", "$runcommand", "$accumulator", "$function"]
        for op in destructive_ops:
            if op in pipeline_str:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Operador no permitido detectado en la agregación: {op}"
                )
                
        if req.tenantId:
            tenant_match = {"tenantId": req.tenantId}
            if req.politicaId and collection_name in ["tramites", "registros_actividad"]:
                tenant_match["politicaId"] = req.politicaId
                
            has_match = False
            for stage in pipeline:
                if "$match" in stage:
                    stage["$match"].update(tenant_match)
                    has_match = True
                    break
            if not has_match:
                pipeline.insert(0, {"$match": tenant_match})
                
        coll = db[collection_name]
        
        def execute_aggregation():
            return list(coll.aggregate(pipeline))
            
        raw_results = await asyncio.to_thread(execute_aggregation)
        
        def json_serializable(data):
            if isinstance(data, list):
                return [json_serializable(item) for item in data]
            elif isinstance(data, dict):
                return {k: json_serializable(v) for k, v in data.items()}
            elif hasattr(data, "isoformat"):
                return data.isoformat()
            elif hasattr(data, "__str__") and data.__class__.__name__ == "ObjectId":
                return str(data)
            else:
                return data
                
        serialized_results = json_serializable(raw_results)
        
        return {
            "collection": collection_name,
            "pipeline": pipeline,
            "results": serialized_results
        }
        
    except HTTPException:
        raise
    except Exception as e:
        import logging
        logging.exception("Error en nl-to-aggregation")
        raise HTTPException(status_code=500, detail=f"Error en el generador de reportes NLP: {str(e)}")

@router.post("/export-report")
async def export_report(req: ExportReportRequest):
    """
    Ejecuta una consulta NLP y la exporta a PDF, Word o Excel.
    """
    if not groq_client:
        raise HTTPException(status_code=400, detail="Groq no configurado.")
        
    from database import get_db
    db = get_db()
    if db is None:
        raise HTTPException(status_code=500, detail="No se pudo conectar a la base de datos.")
        
    try:
        collection_name = req.collection
        pipeline = req.pipeline
        
        if not pipeline and req.query:
            schemas_info = """
Colecciones disponibles en MongoDB:

1. 'registros_actividad': Registro histórico detallado.
   - tramiteId: string
   - estado: string ("HECHO", "EN_PROGRESO")

2. 'tramites': Instancias de procesos de trámites.
   - _id: string
   - politicaId: string
   - tenantId: string
   - estado: string

3. 'clientes': Base de datos de clientes.
   - nombre: string
   - correo: string

4. 'departamentos': Departamentos.
   - nombre: string
   - presupuesto: number
"""
            system_prompt = f"""Eres un Ingeniero de Datos y experto en MongoDB.
Tu tarea es traducir una consulta en lenguaje natural a un pipeline de agregación.
{schemas_info}
Responde estrictamente en JSON:
{{
  "collection": "nombre_de_la_coleccion",
  "pipeline": [ ... ]
}}
"""
            user_content = f"Consulta: \"{req.query}\"\n"
            if req.politicaId:
                user_content += f"Filtrar por politicaId: \"{req.politicaId}\"\n"
            if req.tenantId:
                user_content += f"Filtrar por tenantId: \"{req.tenantId}\"\n"

            response = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content},
                ],
                temperature=0.0,
                response_format={"type": "json_object"}
            )
            result_json = json.loads(response.choices[0].message.content)
            collection_name = result_json.get("collection")
            pipeline = result_json.get("pipeline", [])

        if not collection_name or not pipeline:
            raise HTTPException(status_code=400, detail="Colección o consulta inválida.")

        allowed_collections = ["tramites", "registros_actividad", "clientes", "departamentos", "usuarios", "audit_log", "politicas_negocio"]
        if collection_name not in allowed_collections:
            raise HTTPException(status_code=400, detail="Colección no permitida.")
            
        pipeline_str = json.dumps(pipeline).lower()
        destructive_ops = ["$out", "$merge", "$write", "$eval", "$runcommand", "$accumulator", "$function"]
        for op in destructive_ops:
            if op in pipeline_str:
                raise HTTPException(status_code=400, detail="Operador no permitido detectado.")

        if req.tenantId:
            tenant_match = {"tenantId": req.tenantId}
            if req.politicaId and collection_name in ["tramites", "registros_actividad"]:
                tenant_match["politicaId"] = req.politicaId
                
            has_match = False
            for stage in pipeline:
                if "$match" in stage:
                    stage["$match"].update(tenant_match)
                    has_match = True
                    break
            if not has_match:
                pipeline.insert(0, {"$match": tenant_match})

        coll = db[collection_name]
        
        def execute_aggregation():
            return list(coll.aggregate(pipeline))
            
        raw_results = await asyncio.to_thread(execute_aggregation)
        
        if not raw_results:
            raise HTTPException(
                status_code=400, 
                detail="No se encontraron registros de datos para los criterios seleccionados."
            )

        def json_serializable(data):
            if isinstance(data, list):
                return [json_serializable(item) for item in data]
            elif isinstance(data, dict):
                return {k: json_serializable(v) for k, v in data.items()}
            elif hasattr(data, "isoformat"):
                return data.isoformat()
            elif hasattr(data, "__str__") and data.__class__.__name__ == "ObjectId":
                return str(data)
            else:
                return data
                
        serialized_results = json_serializable(raw_results)

        cleaned_results = []
        for row in serialized_results:
            if not isinstance(row, dict):
                cleaned_results.append(row)
                continue
            cleaned_row = {}
            for k, v in row.items():
                if k in ["_class", "tenantId", "ejecutadoPorId", "proyectoId", "politicaId"]:
                    continue
                if k == "_id":
                    if isinstance(v, str) and len(v) == 24 and all(c in "0123456789abcdefABCDEF" for c in v):
                        continue
                    if isinstance(v, dict) and "$oid" in v:
                        continue
                cleaned_row[k] = v
            cleaned_results.append(cleaned_row)
        serialized_results = cleaned_results

        fmt = req.format.lower()
        if fmt == "xlsx":
            file_stream = generate_excel_report(serialized_results, f"Reporte de {collection_name}")
            filename = f"reporte-{collection_name}-{DateNowStr()}.xlsx"
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif fmt == "docx":
            file_stream = generate_docx_report(serialized_results, f"Reporte de {collection_name}", req.query, req.tenantId, req.politicaId, req.chartImage)
            filename = f"reporte-{collection_name}-{DateNowStr()}.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fmt == "pdf":
            file_stream = generate_pdf_report(serialized_results, f"Reporte de {collection_name}", req.query, req.tenantId, req.politicaId, req.chartImage)
            filename = f"reporte-{collection_name}-{DateNowStr()}.pdf"
            media_type = "application/pdf"
        else:
            raise HTTPException(status_code=400, detail=f"Formato no soportado: {req.format}")

        return StreamingResponse(
            file_stream,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        import logging
        logging.exception("Error en export-report")
        raise HTTPException(status_code=500, detail=f"Error al generar reporte: {str(e)}")

@router.post("/predict-route", response_model=PredictRouteResponse)
async def predict_route(req: PredictRouteRequest):
    if config.ort_routing_session is None or config.ort_anomaly_session is None or config.tf_preprocessing is None:
        raise HTTPException(status_code=503, detail="ONNX models are not loaded.")
        
    try:
        import numpy as np
        import math
        
        def safe_float(val, default=0.0):
            f = float(val)
            if math.isnan(f) or math.isinf(f):
                return default
            return f
        
        scaler = config.tf_preprocessing["scaler"]
        le_depto = config.tf_preprocessing["le_depto"]
        le_politica = config.tf_preprocessing["le_politica"]
        le_actividad = config.tf_preprocessing["le_actividad"]
        
        try:
            depto_idx = le_depto.transform([req.departamento_id])[0]
        except:
            depto_idx = 0
            
        try:
            pol_idx = le_politica.transform([req.politica_id])[0]
        except:
            pol_idx = 0
            
        features = np.array([[
            req.hora_del_dia,
            req.dia_de_semana,
            depto_idx,
            pol_idx,
            req.carga_actual,
            req.historial_cliente
        ]])
        
        features_scaled = scaler.transform(features)
        features_scaled = np.nan_to_num(features_scaled, nan=0.0, posinf=0.0, neginf=0.0)
        
        input_data = features_scaled.astype(np.float32)
        preds = config.ort_routing_session.run(None, {config.ort_routing_input_name: input_data})
        
        duracion_pred = safe_float(preds[0][0][0], 30.0)
        if duracion_pred < 0: duracion_pred = 0.0
            
        prioridad_idx = int(np.argmax(preds[1][0]))
        map_prioridad = {0: "BAJA", 1: "MEDIA", 2: "ALTA"}
        prioridad_str = map_prioridad.get(prioridad_idx, "MEDIA")
        
        ruta_idx = int(np.argmax(preds[2][0]))
        try:
            ruta_str = le_actividad.inverse_transform([ruta_idx])[0]
        except:
            ruta_str = "DEFAULT_NEXT"
            
        reconstruction = config.ort_anomaly_session.run(None, {config.ort_anomaly_input_name: input_data})[0]
        mse = np.mean(np.power(features_scaled - reconstruction, 2), axis=1)
        mse_val = safe_float(mse[0], 0.0)
        is_anomalo = bool(mse_val > 2.5)
        
        score = 1.0 - (duracion_pred / 300.0)
        score = max(0.1, min(0.99, score))
        
        return PredictRouteResponse(
            rutaSugerida=ruta_str,
            tiempoEstimadoMinutos=round(duracion_pred, 1),
            prioridadRecomendada=prioridad_str,
            isAnomalo=is_anomalo,
            scoreEficiencia=round(score, 2)
        )
        
    except Exception as e:
        import logging
        logging.error(f"Error in predict_route: {str(e)}")
        raise HTTPException(status_code=500, detail="Error durante inferencia de ONNX Runtime.")
